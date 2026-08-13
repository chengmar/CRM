from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "dist" / "外贸获客智能体-安装资料获取与使用手册.md"
OUTPUT = ROOT / "dist" / "外贸获客智能体-安装资料获取与使用手册.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
FONT_MONO = "Consolas"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "243447"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "C7D0DA"
RISK = "9B1C1C"
GOLD = "7A5A00"
USABLE_DXA = 9360
TABLE_INDENT_DXA = 120

PAGE_BREAK_HEADINGS: set[str] = set()


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != USABLE_DXA:
        raise ValueError(f"Table widths must total {USABLE_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def set_run_font(run, name=FONT_CN, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN if name == FONT_CN else name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN if name == FONT_CN else name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CN)
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\)|https?://[^\s，。；、]+)")


def add_inline(paragraph, text: str, base_size=11, base_color=INK):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=FONT_MONO, size=max(base_size - 0.5, 8), color=DARK_BLUE)
            run.font.highlight_color = None
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            clean = token.rstrip(".，。；、")
            suffix = token[len(clean) :]
            add_hyperlink(paragraph, clean, clean)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size=base_size, color=base_color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def apply_paragraph_tokens(paragraph, after=6, before=0, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_paragraph(doc, text: str, style=None, align=None, after=6):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    add_inline(paragraph, text)
    apply_paragraph_tokens(paragraph, after=after)
    return paragraph


def add_callout(doc, text: str, kind="info"):
    table = doc.add_table(rows=1, cols=1)
    set_row_cant_split(table.rows[0])
    set_table_geometry(table, [USABLE_DXA])
    set_table_borders(table, color="B8C5D1", size="4")
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=130, bottom=130, start=180, end=180)
    set_cell_shading(cell, "FFF4E5" if kind == "warning" else CALLOUT)
    p = cell.paragraphs[0]
    apply_paragraph_tokens(p, after=0, line=1.2)
    prefix = "重要：" if kind == "warning" else "提示："
    r = p.add_run(prefix)
    set_run_font(r, size=10.5, bold=True, color=RISK if kind == "warning" else DARK_BLUE)
    add_inline(p, text, base_size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc, lines: list[str], language: str):
    table = doc.add_table(rows=1, cols=1)
    set_row_cant_split(table.rows[0])
    set_table_geometry(table, [USABLE_DXA])
    set_table_borders(table, color="D7DCE2", size="4")
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    apply_paragraph_tokens(p, after=0, line=1.08)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name=FONT_MONO, size=8.8, color="263238")
        if idx < len(lines) - 1:
            run.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def choose_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    normalized = "|".join(headers)
    if count == 4 and "安装器字段" in normalized:
        return [1450, 3680, 2860, 1370]
    if count == 3:
        return [1900, 5000, 2460]
    if count == 2:
        return [2300, 7060]
    if count == 4:
        return [1700, 3100, 3100, 1460]
    base = USABLE_DXA // count
    widths = [base] * count
    widths[-1] += USABLE_DXA - sum(widths)
    return widths


def add_markdown_table(doc, headers: list[str], rows: list[list[str]]):
    widths = choose_widths(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    table.rows[0]._tr.get_or_add_trPr()
    font_size = 8.2 if len(headers) == 4 else 9.0

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_paragraph_tokens(p, after=0, line=1.05)
        add_inline(p, text, base_size=font_size, base_color=DARK_BLUE)
        for run in p.runs:
            run.bold = True

    for row_index, values in enumerate(rows):
        new_row = table.add_row()
        set_row_cant_split(new_row)
        row_cells = new_row.cells
        for idx in range(len(headers)):
            cell = row_cells[idx]
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == len(headers) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            apply_paragraph_tokens(p, after=0, line=1.08)
            add_inline(p, values[idx] if idx < len(values) else "", base_size=font_size)

    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AEB8C4")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    p.paragraph_format.space_after = Pt(8)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=8.5, color=MUTED)


def create_decimal_abstract_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CN)
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def create_numbering_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_ref])


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Cover Subtitle" not in styles:
        cover_subtitle = styles.add_style("Cover Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cover_subtitle = styles["Cover Subtitle"]
    cover_subtitle.font.name = FONT_CN
    cover_subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    cover_subtitle.font.size = Pt(13)
    cover_subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    cover_subtitle.paragraph_format.space_after = Pt(8)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("外贸获客智能体")
    set_run_font(r, size=30, bold=True, color=DARK_BLUE)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(18)
    r = p2.add_run("安装资料获取与使用手册")
    set_run_font(r, size=22, bold=True, color=BLUE)

    add_horizontal_rule(doc)

    subtitle = doc.add_paragraph(style="Cover Subtitle")
    add_inline(subtitle, "给外贸老板的逐项填写说明，不需要 GitHub、Docker、命令行或编程知识。", base_size=13, base_color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2300, 7060])
    set_table_borders(meta, color="D5DCE4", size="4")
    values = [
        ("配套版本", "免费测试版一键安装器 0.2.2"),
        ("适用对象", "外贸老板、销售负责人、安装协助人员"),
        ("更新日期", "2026-07-17"),
        ("核心原则", "真实资料、最小权限、人工审批、询价后人工接管"),
    ]
    for row_idx, (label, value) in enumerate(values):
        for col_idx, text in enumerate((label, value)):
            cell = meta.cell(row_idx, col_idx)
            set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
            if col_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            apply_paragraph_tokens(p, after=0, line=1.1)
            add_inline(p, text, base_size=10.5, base_color=DARK_BLUE if col_idx == 0 else INK)
            if col_idx == 0:
                for run in p.runs:
                    run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_callout(
        doc,
        "安装成功后 Gmail 具备真实外发能力但保持暂停；必须先自发测试并在飞书确认启用；WhatsApp 测试阶段保持关闭。",
        kind="warning",
    )

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(24)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p3.add_run("配套安装文件")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p4 = doc.add_paragraph()
    add_inline(p4, "外贸获客智能体-免费测试版-一键安装.exe", base_size=10.5)
    apply_paragraph_tokens(p4, after=2)
    p5 = doc.add_paragraph()
    add_inline(p5, "SHA-256: 1eef6cc57f475401daad39df04a58d009ec671e65a1fa55adc99e4886f3e696f", base_size=8.8, base_color=MUTED)
    apply_paragraph_tokens(p5, after=0)

    doc.add_page_break()


def parse_table(lines: list[str], start: int):
    def split(row: str) -> list[str]:
        return [cell.strip() for cell in row.strip().strip("|").split("|")]

    headers = split(lines[start])
    rows = []
    index = start + 2
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        rows.append(split(lines[index]))
        index += 1
    return headers, rows, index


def parse_body(doc: Document, markdown: str):
    body = markdown.split("<!-- BODY -->", 1)[1]
    lines = body.splitlines()
    index = 0
    paragraph_buffer: list[str] = []
    decimal_abstract_id = create_decimal_abstract_numbering(doc)
    active_num_id: int | None = None

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(part.strip().rstrip("  ") for part in paragraph_buffer).strip()
            if text:
                add_paragraph(doc, text)
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        is_numbered = bool(re.match(r"^\d+\.\s+", stripped))

        if not stripped:
            flush_paragraph()
            active_num_id = None
            index += 1
            continue

        if not is_numbered:
            active_num_id = None

        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines, language)
            index += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            title = stripped[3:].strip()
            if title in PAGE_BREAK_HEADINGS and len(doc.paragraphs) > 3:
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, title, base_size=16, base_color=BLUE)
            index += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:].strip(), base_size=13, base_color=BLUE)
            index += 1
            continue

        if stripped.startswith("#### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:].strip(), base_size=12, base_color=DARK_BLUE)
            index += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            quote = stripped[2:].strip()
            kind = "warning" if quote.startswith("重要：") else "info"
            quote = re.sub(r"^(重要|提示)：", "", quote).strip()
            add_callout(doc, quote, kind=kind)
            index += 1
            continue

        if stripped == "---":
            flush_paragraph()
            add_horizontal_rule(doc)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[index + 1]):
            flush_paragraph()
            headers, rows, index = parse_table(lines, index)
            add_markdown_table(doc, headers, rows)
            continue

        if re.match(r"^-\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^-\s+", "", stripped))
            apply_paragraph_tokens(p, after=4)
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            if active_num_id is None:
                active_num_id = create_numbering_instance(doc, decimal_abstract_id)
            p = doc.add_paragraph()
            apply_numbering(p, active_num_id)
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            apply_paragraph_tokens(p, after=4)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def configure_headers_footers(doc: Document):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D4DAE1")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    r = p.add_run("外贸获客智能体  |  安装资料获取与使用手册")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def audit_tables(doc: Document):
    for idx, table in enumerate(doc.tables, start=1):
        grid = table._tbl.tblGrid
        widths = [int(col.get(qn("w:w"))) for col in grid]
        if sum(widths) != USABLE_DXA:
            raise RuntimeError(f"Table {idx} grid total is {sum(widths)}, expected {USABLE_DXA}")
        for row in table.rows:
            if len(row.cells) != len(widths):
                continue
            for cell_idx, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None or int(tc_w.get(qn("w:w"))) != widths[cell_idx]:
                    raise RuntimeError(f"Table {idx} cell width mismatch")


def main() -> int:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_section(doc.sections[0])
    configure_styles(doc)
    configure_headers_footers(doc)
    add_cover(doc)
    parse_body(doc, markdown)

    core = doc.core_properties
    core.title = "外贸获客智能体：安装资料获取与使用手册"
    core.subject = "一键安装器资料准备、账号配置、安装和日常使用"
    core.author = "CRM Agent"
    core.keywords = "外贸获客, AI Agent, 飞书, Gmail, Ubuntu, 安装手册"
    core.comments = "不含真实密码、Token 或客户私有资料。"

    audit_tables(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
